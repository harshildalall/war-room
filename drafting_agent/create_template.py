from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

doc.add_paragraph('{{ date_today }}')
doc.add_paragraph('')
doc.add_paragraph('Re: Appeal of Denied Claim — Case {{ case_id }}')
doc.add_paragraph('')
doc.add_paragraph('To Whom It May Concern,')
doc.add_paragraph('')
doc.add_paragraph('{{ letter_body }}')
doc.add_paragraph('')
doc.add_paragraph('Recommended Remedy: {{ recommended_remedy }}')
doc.add_paragraph('Appeal Strength: {{ confidence_score }}')
doc.add_paragraph('')
doc.add_paragraph('Key Arguments:')
doc.add_paragraph('{% for arg in strongest_arguments %}- {{ arg }}{% endfor %}')
doc.add_paragraph('')
doc.add_paragraph('Contract Violations:')
doc.add_paragraph('{% for v in contract_violations %}- {{ v.clause }} (Score: {{ v.contradiction_score }}){% endfor %}')
doc.add_paragraph('')
doc.add_paragraph('Footnotes:')
doc.add_paragraph('{% for fn in citations_footnoted %}[{{ fn.footnote_index }}] {{ fn.source }} — {{ fn.quote }}{% endfor %}')
doc.add_paragraph('')
doc.add_paragraph('Exhibits:')
doc.add_paragraph('{% for ex in exhibits_checklist %}{{ ex.exhibit_label }}: {{ ex.description }}{% endfor %}')
doc.add_paragraph('')
doc.add_paragraph('Submission Steps:')
doc.add_paragraph('{% for step in submission_instructions %}{{ loop.index }}. {{ step }}{% endfor %}')
doc.add_paragraph('')
doc.add_paragraph('Deadline: {{ deadline }}')
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('')
doc.add_paragraph('_______________________')
doc.add_paragraph('[Patient Name]')

import os
os.makedirs('templates', exist_ok=True)
doc.save('templates/appeal_letter_tpl.docx')
print('Template created at templates/appeal_letter_tpl.docx')
